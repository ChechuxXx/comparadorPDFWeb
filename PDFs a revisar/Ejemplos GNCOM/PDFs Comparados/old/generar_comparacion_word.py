"""
Script para generar documento Word con comparación visual de PDFs
FAC_001_ES_PI25142000355153 - Papyrus vs Servinform

Requisitos:
pip install PyMuPDF python-docx Pillow

Autor: Análisis automatizado
Fecha: 12/01/2026
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os
from datetime import datetime

# Configuración de rutas
BASE_DIR = r"C:\Users\GONZALOGUERRASABUGO\Box\BAP COE Delta\DAZEL - Papyrus - Solución conceptual\PDFs a revisar\Ejemplos GNCOM"
PAPYRUS_PDF = os.path.join(BASE_DIR, "Papyrus", "Español", "FAC_001_ES_PI25142000355153.pdf")
SERVINFORM_PDF = os.path.join(BASE_DIR, "Servinform", "Español", "FAC_001_ES_PI25142000355153.pdf")
OUTPUT_DIR = os.path.join(BASE_DIR, "PDFs Comparados")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "FAC_001_Comparacion_Visual.docx")
TEMP_DIR = os.path.join(OUTPUT_DIR, "temp_images")

# Crear directorio temporal para imágenes
os.makedirs(TEMP_DIR, exist_ok=True)

def extract_pdf_page_as_image(pdf_path, page_num, output_path, zoom=2.0):
    """
    Extrae una página de PDF como imagen
    
    Args:
        pdf_path: Ruta al archivo PDF
        page_num: Número de página (0-indexed)
        output_path: Ruta donde guardar la imagen
        zoom: Factor de zoom para calidad (2.0 = 200%)
    """
    try:
        doc = fitz.open(pdf_path)
        page = doc[page_num]
        
        # Matriz de transformación para zoom
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        
        # Guardar imagen
        pix.save(output_path)
        doc.close()
        
        return True
    except Exception as e:
        print(f"Error extrayendo página {page_num} de {pdf_path}: {e}")
        return False

def extract_pdf_region(pdf_path, page_num, rect, output_path, zoom=2.0):
    """
    Extrae una región específica de una página PDF
    
    Args:
        pdf_path: Ruta al archivo PDF
        page_num: Número de página (0-indexed)
        rect: Tupla (x0, y0, x1, y1) con coordenadas de la región
        output_path: Ruta donde guardar la imagen
        zoom: Factor de zoom
    """
    try:
        doc = fitz.open(pdf_path)
        page = doc[page_num]
        
        # Crear rectángulo
        clip_rect = fitz.Rect(rect)
        
        # Matriz de transformación
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, clip=clip_rect)
        
        pix.save(output_path)
        doc.close()
        
        return True
    except Exception as e:
        print(f"Error extrayendo región de {pdf_path}: {e}")
        return False

def add_heading(doc, text, level=1):
    """Añade un encabezado con formato"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_difference_section(doc, num, title, papyrus_desc, servinform_desc, 
                          papyrus_img=None, servinform_img=None):
    """
    Añade una sección de diferencia al documento
    
    Args:
        doc: Documento de Word
        num: Número de diferencia
        title: Título de la diferencia
        papyrus_desc: Descripción de Papyrus
        servinform_desc: Descripción de Servinform
        papyrus_img: Ruta a imagen de Papyrus (opcional)
        servinform_img: Ruta a imagen de Servinform (opcional)
    """
    # Título de la diferencia
    heading = doc.add_heading(f"{num}. {title}", level=2)
    
    # Crear tabla para comparación lado a lado
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'PAPYRUS'
    hdr_cells[1].text = 'SERVINFORM'
    
    # Hacer encabezados en negrita
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Descripciones
    desc_cells = table.rows[1].cells
    desc_cells[0].text = papyrus_desc
    desc_cells[1].text = servinform_desc
    
    # Imágenes (si existen)
    img_cells = table.rows[2].cells
    
    if papyrus_img and os.path.exists(papyrus_img):
        try:
            paragraph = img_cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(papyrus_img, width=Inches(3.0))
        except Exception as e:
            img_cells[0].text = f"[Error cargando imagen: {e}]"
    else:
        img_cells[0].text = "[Imagen no disponible]"
    
    if servinform_img and os.path.exists(servinform_img):
        try:
            paragraph = img_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(servinform_img, width=Inches(3.0))
        except Exception as e:
            img_cells[1].text = f"[Error cargando imagen: {e}]"
    else:
        img_cells[1].text = "[Imagen no disponible]"
    
    # Espacio después de la tabla
    doc.add_paragraph()

def create_comparison_document():
    """Función principal para crear el documento de comparación"""
    
    print("Iniciando generación de documento de comparación...")
    print(f"Papyrus PDF: {PAPYRUS_PDF}")
    print(f"Servinform PDF: {SERVINFORM_PDF}")
    
    # Verificar que los PDFs existen
    if not os.path.exists(PAPYRUS_PDF):
        print(f"ERROR: No se encuentra el PDF de Papyrus: {PAPYRUS_PDF}")
        return False
    
    if not os.path.exists(SERVINFORM_PDF):
        print(f"ERROR: No se encuentra el PDF de Servinform: {SERVINFORM_PDF}")
        return False
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading('COMPARACIÓN VISUAL DE DOCUMENTOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('FAC_001_ES_PI25142000355153', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    doc.add_paragraph(f"Fecha de análisis: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(f"Papyrus: Papyrus/Español/FAC_001_ES_PI25142000355153.pdf")
    doc.add_paragraph(f"Servinform: Servinform/Español/FAC_001_ES_PI25142000355153.pdf")
    doc.add_paragraph()
    
    # Resumen ejecutivo
    add_heading(doc, "RESUMEN EJECUTIVO", level=1)
    
    p = doc.add_paragraph()
    p.add_run("DIFERENCIA CRÍTICA: ").bold = True
    p.add_run("La tasa GTS difiere entre ambos documentos (1,325% en Papyrus vs 1,354% en Servinform), ")
    run = p.add_run("lo cual podría afectar cálculos financieros.")
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph("Se han identificado 11 diferencias de contenido y 4 diferencias de formato/tipografía.")
    doc.add_paragraph()
    
    # Extraer imágenes de páginas completas para referencia
    print("\nExtrayendo imágenes de páginas completas...")
    
    papyrus_page1 = os.path.join(TEMP_DIR, "papyrus_page1.png")
    servinform_page1 = os.path.join(TEMP_DIR, "servinform_page1.png")
    
    extract_pdf_page_as_image(PAPYRUS_PDF, 0, papyrus_page1, zoom=1.5)
    extract_pdf_page_as_image(SERVINFORM_PDF, 0, servinform_page1, zoom=1.5)
    
    # Sección de diferencias
    add_heading(doc, "DIFERENCIAS DETALLADAS", level=1)
    
    print("\nGenerando secciones de diferencias...")
    
    # Diferencia 1: Información administrativa
    print("  - Diferencia 1: Información administrativa")
    add_difference_section(
        doc, 1,
        "INFORMACIÓN DE CABECERA ADMINISTRATIVA",
        "Muestra la información administrativa (UNIDAD TRAMITADORA, OFICINA CONTABLE, ÓRGANO GESTOR) al INICIO del documento",
        "Muestra esta misma información al FINAL de la primera página",
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 2: Dirección en datos sociales
    print("  - Diferencia 2: Dirección en datos sociales")
    add_difference_section(
        doc, 2,
        'DIRECCIÓN EN "DATOS SOCIALES"',
        '"C EDUARDO GARCIA S/NBJ" (sin espacio antes de BJ)',
        '"C EDUARDO GARCIA S/N" (sin el "BJ")',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 3: Dirección en domicilio de envío
    print("  - Diferencia 3: Dirección en domicilio de envío")
    add_difference_section(
        doc, 3,
        'DIRECCIÓN EN "DOMICILIO DE ENVIO"',
        '"C EDUARDO GARCIA S/NBJ" y "39011- PEÑACASTILLO" (con guion pegado)',
        '"C EDUARDO GARCIA S/N  BJ" (con dos espacios antes de BJ) y "39011 SANTANDER" (sin guion, sin PEÑACASTILLO)',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 4: Formato de conceptos
    print("  - Diferencia 4: Formato de conceptos")
    add_difference_section(
        doc, 4,
        "FORMATO DE CONCEPTOS DE FACTURACIÓN",
        'Los conceptos aparecen sin espacios (ej: "CUOTAFIJA", "TÉRMINOVARIABLE")',
        'Los conceptos tienen espacios (ej: "CUOTA FIJA", "TÉRMINO VARIABLE")',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 5: Unidades monetarias
    print("  - Diferencia 5: Unidades monetarias")
    add_difference_section(
        doc, 5,
        "UNIDADES MONETARIAS",
        '"Eur" (con mayúscula inicial)',
        '"EUR" (todo en mayúsculas)',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 6: Información de contacto
    print("  - Diferencia 6: Información de contacto")
    add_difference_section(
        doc, 6,
        "INFORMACIÓN DE CONTACTO",
        'Incluye "GNCom(Att.Reclamaciones)Avenida de San Luis 77, 28033 Madrid"',
        'Incluye "Dirección Postal Avenida de San Luis 77, 28033 Madrid" (sin mención a reclamaciones)',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 7: Pie de página
    print("  - Diferencia 7: Pie de página")
    add_difference_section(
        doc, 7,
        "PIE DE PÁGINA PRIMERA HOJA",
        'Incluye código largo "0142/99/F001/E/25329/1/F031/V3.04.10/DTGCFT60/1764067788505-1-1" y "1 de 3"',
        'No incluye este código, solo muestra "ef 01100001"',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 8: Marca de agua
    print("  - Diferencia 8: Marca de agua")
    add_difference_section(
        doc, 8,
        "MARCA DE AGUA/ETIQUETA",
        'Muestra "DUPLICADO" en varias páginas con "Documento informativo sin valor legal generado a partir de datos factura-e"',
        'No muestra esta marca de "DUPLICADO"',
        papyrus_page1,
        servinform_page1
    )
    
    # Diferencia 9: TASA GTS (CRÍTICA)
    print("  - Diferencia 9: TASA GTS *** CRÍTICA ***")
    
    # Extraer páginas 2 y 3 para mostrar la tasa GTS
    papyrus_page3 = os.path.join(TEMP_DIR, "papyrus_page3.png")
    servinform_page3 = os.path.join(TEMP_DIR, "servinform_page3.png")
    
    extract_pdf_page_as_image(PAPYRUS_PDF, 2, papyrus_page3, zoom=1.5)
    extract_pdf_page_as_image(SERVINFORM_PDF, 2, servinform_page3, zoom=1.5)
    
    # Añadir sección con énfasis
    heading = doc.add_heading("9. TASA GTS *** DIFERENCIA NUMÉRICA IMPORTANTE ***", level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'PAPYRUS'
    hdr_cells[1].text = 'SERVINFORM'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 0, 0)
    
    desc_cells = table.rows[1].cells
    desc_cells[0].text = '"GTS(1,325%)"'
    desc_cells[1].text = '"GTS(1,354%)"'
    
    img_cells = table.rows[2].cells
    
    if os.path.exists(papyrus_page3):
        paragraph = img_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(papyrus_page3, width=Inches(3.0))
    
    if os.path.exists(servinform_page3):
        paragraph = img_cells[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(servinform_page3, width=Inches(3.0))
    
    doc.add_paragraph()
    
    # Diferencia 10: Numeración de páginas
    print("  - Diferencia 10: Numeración de páginas")
    papyrus_page2 = os.path.join(TEMP_DIR, "papyrus_page2.png")
    servinform_page2 = os.path.join(TEMP_DIR, "servinform_page2.png")
    
    extract_pdf_page_as_image(PAPYRUS_PDF, 1, papyrus_page2, zoom=1.5)
    extract_pdf_page_as_image(SERVINFORM_PDF, 1, servinform_page2, zoom=1.5)
    
    add_difference_section(
        doc, 10,
        "NUMERACIÓN DE PÁGINAS",
        '"2 de 3", "3 de 3"',
        '"ef 02000001" (código diferente)',
        papyrus_page2,
        servinform_page2
    )
    
    # Diferencia 11: Gráfico de consumo
    print("  - Diferencia 11: Gráfico de consumo histórico")
    add_difference_section(
        doc, 11,
        "GRÁFICO DE CONSUMO HISTÓRICO",
        'Muestra valor "0,801" en el eje',
        'Muestra "00" en el mismo lugar',
        papyrus_page2,
        servinform_page2
    )
    
    # Sección de conclusiones
    add_heading(doc, "CONCLUSIONES Y RECOMENDACIONES", level=1)
    
    doc.add_paragraph("Basándose en el análisis visual comparativo:")
    
    conclusions = [
        "La diferencia más crítica es la tasa GTS (1,325% vs 1,354%) que debe ser verificada y corregida.",
        "Las diferencias de formato (espaciado, mayúsculas) son consistentes y sugieren diferentes motores de renderizado.",
        "La ubicación de información administrativa difiere, lo que puede afectar la experiencia del usuario.",
        "Papyrus incluye más información de trazabilidad (códigos largos, marca DUPLICADO).",
        "Servinform tiene mejor legibilidad general debido al espaciado correcto.",
        "Se recomienda estandarizar el formato de direcciones y conceptos.",
        "Es necesario definir cuál es la tasa GTS correcta aplicable."
    ]
    
    for conclusion in conclusions:
        p = doc.add_paragraph(conclusion, style='List Bullet')
    
    # Guardar documento
    print(f"\nGuardando documento en: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    
    print(f"\n✓ Documento generado exitosamente: {OUTPUT_DOCX}")
    print(f"✓ Imágenes temporales guardadas en: {TEMP_DIR}")
    
    return True

if __name__ == "__main__":
    try:
        success = create_comparison_document()
        if success:
            print("\n" + "="*60)
            print("PROCESO COMPLETADO EXITOSAMENTE")
            print("="*60)
            print(f"\nDocumento Word generado: {OUTPUT_DOCX}")
            print("\nPuedes eliminar la carpeta 'temp_images' si no necesitas las imágenes individuales.")
        else:
            print("\n" + "="*60)
            print("ERROR: El proceso no se completó correctamente")
            print("="*60)
    except Exception as e:
        print(f"\nERROR FATAL: {e}")
        import traceback
        traceback.print_exc()