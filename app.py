"""
Comparador PDF Web Application v2.0
Versión Web - Basado en v2.9.1
Desarrollado por: Jesus Eduardo Soler Collantes
"""

from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import os
import uuid
import fitz
from PIL import Image, ImageDraw, ImageFont
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import threading
from werkzeug.utils import secure_filename
import shutil

app = Flask(__name__)
app.secret_key = 'comparador-pdf-secret-key-2026-v2'
CORS(app)

# Configuración
UPLOAD_FOLDER = 'uploads'
RESULTS_FOLDER = 'results'
ALLOWED_EXTENSIONS = {'pdf'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(RESULTS_FOLDER, exist_ok=True)

# Almacenamiento de progreso
comparison_progress = {}
batch_progress = {}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', str(text))

def normalize_text(text):
    text = ' '.join(text.split()).strip()
    text = text.replace(' %', '%').replace(' €', '€').replace(' ,', ',').replace(' .', '.')
    text = text.replace('( ', '(').replace(' )', ')')
    return text

def capture_area(pdf_path, page_num, word, temp_dir, label, zoom=5.0, mark_missing=False, ref_rect=None):
    """
    Captura un área del PDF con el texto resaltado.
    Versión web con búsqueda mejorada de frases.
    """
    doc = fitz.open(pdf_path)
    page = doc[page_num]
    
    zoom_page = 2.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom_page, zoom_page))
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    draw = ImageDraw.Draw(img)
    
    # Estrategia 1: Búsqueda exacta con preservación de espacios
    instances = page.search_for(word, flags=fitz.TEXT_PRESERVE_WHITESPACE)
    
    # Estrategia 2: Búsqueda sin flags especiales
    if not instances:
        instances = page.search_for(word)
    
    # Estrategia 3: Búsqueda palabra por palabra si es una frase (CON FILTRO DE PROXIMIDAD)
    if not instances and ' ' in word:
        words = word.split()
        
        # Buscar todas las instancias de cada palabra
        words_instances = []
        for w in words:
            word_rects = page.search_for(w, flags=fitz.TEXT_PRESERVE_WHITESPACE)
            if not word_rects:
                word_rects = page.search_for(w)
            if word_rects:
                words_instances.append((w, word_rects))
        
        # Si encontramos todas las palabras, buscar grupos cercanos
        if len(words_instances) == len(words):
            # Buscar la primera palabra
            first_word_rects = words_instances[0][1]
            
            best_match = None
            best_distance = float('inf')
            
            # Para cada instancia de la primera palabra
            for first_rect in first_word_rects:
                candidate_rects = [first_rect]
                total_distance = 0
                valid_sequence = True
                
                # Intentar encontrar las siguientes palabras cerca
                last_rect = first_rect
                for i in range(1, len(words_instances)):
                    word_name, word_rects = words_instances[i]
                    
                    # Buscar la instancia más cercana de esta palabra
                    closest_rect = None
                    min_dist = float('inf')
                    
                    for rect in word_rects:
                        # Calcular distancia horizontal (debe estar a la derecha y cerca)
                        horizontal_dist = rect.x0 - last_rect.x1
                        vertical_dist = abs(rect.y0 - last_rect.y0)
                        
                        # Solo considerar si está en la misma línea (vertical < 5) y cerca (horizontal < 100)
                        if vertical_dist < 5 and 0 <= horizontal_dist < 100:
                            dist = horizontal_dist + vertical_dist
                            if dist < min_dist:
                                min_dist = dist
                                closest_rect = rect
                    
                    if closest_rect:
                        candidate_rects.append(closest_rect)
                        total_distance += min_dist
                        last_rect = closest_rect
                    else:
                        valid_sequence = False
                        break
                
                # Si encontramos todas las palabras en secuencia cercana
                if valid_sequence and len(candidate_rects) == len(words):
                    if total_distance < best_distance:
                        best_distance = total_distance
                        best_match = candidate_rects
            
            # Si encontramos una secuencia válida, crear rectángulo que la englobe
            if best_match:
                min_x0 = min(r.x0 for r in best_match)
                min_y0 = min(r.y0 for r in best_match)
                max_x1 = max(r.x1 for r in best_match)
                max_y1 = max(r.y1 for r in best_match)
                instances = [fitz.Rect(min_x0, min_y0, max_x1, max_y1)]
    
    # Estrategia 4: Búsqueda case-insensitive (solo si las anteriores fallan)
    if not instances:
        page_text = page.get_text("text").lower()
        if word.lower() in page_text:
            # Intentar búsqueda con variaciones de mayúsculas/minúsculas
            for search_term in [word, word.lower(), word.upper(), word.capitalize()]:
                instances = page.search_for(search_term)
                if instances:
                    break
    
    if instances:
        # TEXTO ENCONTRADO - Resaltar en VERDE
        for rect in instances:
            x0s = rect.x0 * zoom_page
            y0s = rect.y0 * zoom_page
            x1s = rect.x1 * zoom_page
            y1s = rect.y1 * zoom_page
            for i in range(3):
                draw.ellipse([x0s - 10 - i, y0s - 10 - i, x1s + 10 + i, y1s + 10 + i], outline='green')
        
        try:
            font = ImageFont.truetype("arial.ttf", 40)
        except:
            font = ImageFont.load_default()
        draw.text((10, 10), f"Encontrado: {len(instances)} ocurrencias", fill='green', font=font)
    else:
        # TEXTO NO ENCONTRADO
        if mark_missing and ref_rect:
            # Usar posición de referencia del otro PDF
            x0s = ref_rect.x0 * zoom_page
            y0s = ref_rect.y0 * zoom_page
            x1s = ref_rect.x1 * zoom_page
            y1s = ref_rect.y1 * zoom_page
            
            # Dibujar rectángulo ROJO indicando donde DEBERÍA estar
            for i in range(5):
                draw.rectangle([x0s - 15 - i, y0s - 15 - i, x1s + 15 + i, y1s + 15 + i], outline='red')
            
            # Agregar texto explicativo
            try:
                font = ImageFont.truetype("arial.ttf", 40)
            except:
                font = ImageFont.load_default()
            draw.text((10, 10), "TEXTO NO ENCONTRADO", fill='red', font=font)
            draw.text((10, 60), "(Deberia estar aqui segun PDF de referencia)", fill='red', font=font)
        else:
            # Sin referencia, solo indicar que no se encontró
            try:
                font = ImageFont.truetype("arial.ttf", 40)
            except:
                font = ImageFont.load_default()
            draw.text((10, 10), "Texto no encontrado en este PDF", fill='red', font=font)
    
    word_safe = clean_text(word)[:20].replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    filepath = os.path.join(temp_dir, f"pag{page_num+1}_{label}_{word_safe}_FULL.png")
    img.save(filepath)
    doc.close()
    
    # Retornar filepath y el primer rectángulo encontrado (o None)
    return filepath, instances[0] if instances else None

def generate_comparison(pap_pdf, serv_pdf, ref_page_num, comp_page_num, diff, temp_dir):
    """
    Genera comparación de una diferencia entre dos PDFs con capturas de pantalla.
    """
    palabra, tipo = diff['palabra'], diff['tipo']
    
    if tipo == 'SOLO_REFERENCIA':
        # Texto solo en PDF Referencia
        # 1. Capturar en PDF Referencia (con texto resaltado en verde)
        img_p, rect_p = capture_area(pap_pdf, ref_page_num, palabra, temp_dir, "referencia", zoom=5.0)
        # 2. Capturar en PDF a Comparar (marcando donde DEBERÍA estar en rojo)
        img_s, _ = capture_area(serv_pdf, comp_page_num, palabra, temp_dir, "comparar_FALTA", zoom=5.0, mark_missing=True, ref_rect=rect_p)
    else:
        # Texto solo en PDF a Comparar
        # 1. Capturar en PDF a Comparar (con texto resaltado en verde)
        img_s, rect_s = capture_area(serv_pdf, comp_page_num, palabra, temp_dir, "comparar", zoom=5.0)
        # 2. Capturar en PDF Referencia (marcando donde DEBERÍA estar en rojo)
        img_p, _ = capture_area(pap_pdf, ref_page_num, palabra, temp_dir, "referencia_FALTA", zoom=5.0, mark_missing=True, ref_rect=rect_s)
    
    return {'palabra': palabra, 'tipo': tipo, 'img_p': img_p, 'img_s': img_s}

def compare_content(text1, text2, max_phrase_length=10):
    """Compara el contenido de dos textos"""
    t1, t2 = normalize_text(text1), normalize_text(text2)
    if t1.lower() == t2.lower():
        return "IDENTICO", [], "IDENTICO"
    
    w1, w2 = t1.split(), t2.split()
    s1, s2 = set(w.lower() for w in w1), set(w.lower() for w in w2)
    
    solo1, solo2 = s1 - s2, s2 - s1
    
    if solo1 or solo2:
        diffs = []
        
        def agrupar_frases(palabras_faltantes, texto_original, max_words=10):
            if not palabras_faltantes:
                return []
            
            palabras_orig = texto_original.lower().split()
            frases_agrupadas = []
            palabras_procesadas = set()
            palabras_faltantes_lower = set(p.lower() for p in palabras_faltantes)
            
            i = 0
            while i < len(palabras_orig):
                if palabras_orig[i] in palabras_faltantes_lower and palabras_orig[i] not in palabras_procesadas:
                    inicio = i
                    fin = i
                    
                    while fin < len(palabras_orig) - 1:
                        siguiente = fin + 1
                        if palabras_orig[siguiente] in palabras_faltantes_lower:
                            fin = siguiente
                        else:
                            if (siguiente < len(palabras_orig) - 1 and 
                                palabras_orig[siguiente] in ['y', 'de', 'en', 'a', 'el', 'la', 'los', 'las', 'del', 'al', 'con', 'por', 'para'] and
                                palabras_orig[siguiente + 1] in palabras_faltantes_lower):
                                fin = siguiente + 1
                            else:
                                break
                    
                    palabras_frase = palabras_orig[inicio:fin + 1]
                    
                    if len(palabras_frase) > max_words:
                        palabras_frase = palabras_frase[:max_words]
                        fin = inicio + max_words - 1
                    
                    frase = ' '.join(palabras_frase)
                    palabras_en_frase = set(palabras_frase)
                    
                    if palabras_en_frase & palabras_faltantes_lower:
                        frases_agrupadas.append({
                            'frase': frase,
                            'palabras': palabras_en_frase & palabras_faltantes_lower
                        })
                        palabras_procesadas.update(palabras_en_frase & palabras_faltantes_lower)
                    
                    i = fin + 1
                else:
                    i += 1
            
            for palabra in palabras_faltantes_lower:
                if palabra not in palabras_procesadas:
                    frases_agrupadas.append({
                        'frase': palabra,
                        'palabras': {palabra}
                    })
                    palabras_procesadas.add(palabra)
            
            return frases_agrupadas
        
        frases_ref = agrupar_frases(solo1, t1, max_words=max_phrase_length) if solo1 else []
        frases_comp = agrupar_frases(solo2, t2, max_words=max_phrase_length) if solo2 else []
        
        palabras_procesadas_global = set()
        
        for frase_info in frases_ref:
            palabras_frase = frase_info['palabras']
            if not (palabras_frase & palabras_procesadas_global):
                diffs.append({
                    'tipo': 'SOLO_REFERENCIA',
                    'palabra': frase_info['frase']
                })
                palabras_procesadas_global.update(palabras_frase)
        
        for frase_info in frases_comp:
            palabras_frase = frase_info['palabras']
            if not (palabras_frase & palabras_procesadas_global):
                diffs.append({
                    'tipo': 'SOLO_COMPARAR',
                    'palabra': frase_info['frase']
                })
                palabras_procesadas_global.update(palabras_frase)
        
        return "DIFERENTE", diffs[:10], "CONTENIDO"
    
    return "IDENTICO", [], "FORMATO"

def get_pdf_pages(pdf_path):
    """Obtiene el número de páginas de un PDF"""
    try:
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        doc.close()
        return num_pages
    except:
        return 0

def parse_custom_pages(pages_list, total_pages):
    """Convierte lista de páginas a índices (0-based)"""
    if not pages_list:
        return list(range(total_pages))
    
    # pages_list ya viene como lista de números de página (1-based)
    # Convertir a índices (0-based) y filtrar páginas válidas
    indices = []
    for page_num in pages_list:
        if 1 <= page_num <= total_pages:
            indices.append(page_num - 1)
    
    return sorted(set(indices))

def process_comparison_web(task_id, pdf1_path, pdf2_path, output_dir, start_page_ref=1, end_page_ref=None,
                          start_page_comp=1, end_page_comp=None, custom_pages_ref=None, 
                          custom_pages_comp=None, max_errors=500, max_phrase_length=10):
    """Procesa la comparación de PDFs para la versión web con rangos separados y capturas de pantalla"""
    try:
        comparison_progress[task_id] = {
            'status': 'processing',
            'progress': 0,
            'message': 'Iniciando comparación...',
            'errors_content': 0,
            'errors_format': 0,
            'pages_processed': 0,
            'total_pages': 0
        }
        
        # Crear directorio temporal para capturas
        temp_dir = os.path.join(output_dir, f"temp_screenshots_{task_id}")
        os.makedirs(temp_dir, exist_ok=True)
        
        doc1 = fitz.open(pdf1_path)
        doc2 = fitz.open(pdf2_path)
        
        total_pages_doc1 = len(doc1)
        total_pages_doc2 = len(doc2)
        
        # Determinar qué páginas comparar para cada PDF
        if custom_pages_ref and custom_pages_comp:
            # Modo páginas personalizadas - rangos separados
            pages_ref = parse_custom_pages(custom_pages_ref, total_pages_doc1)
            pages_comp = parse_custom_pages(custom_pages_comp, total_pages_doc2)
        elif end_page_ref is None and end_page_comp is None:
            # Modo todas las páginas
            max_pages = min(total_pages_doc1, total_pages_doc2)
            pages_ref = list(range(max_pages))
            pages_comp = list(range(max_pages))
        else:
            # Modo rango - rangos separados
            start_idx_ref = max(0, start_page_ref - 1)
            end_idx_ref = min(total_pages_doc1, end_page_ref) if end_page_ref else total_pages_doc1
            pages_ref = list(range(start_idx_ref, end_idx_ref))
            
            start_idx_comp = max(0, start_page_comp - 1)
            end_idx_comp = min(total_pages_doc2, end_page_comp) if end_page_comp else total_pages_doc2
            pages_comp = list(range(start_idx_comp, end_idx_comp))
        
        # Comparar hasta el mínimo número de páginas disponibles
        num_pages = min(len(pages_ref), len(pages_comp))
        comparison_progress[task_id]['total_pages'] = num_pages
        
        errores_contenido = []
        errores_formato = []
        
        for idx in range(num_pages):
            if comparison_progress[task_id].get('cancelled', False):
                doc1.close()
                doc2.close()
                comparison_progress[task_id]['status'] = 'cancelled'
                return
            
            page_idx_ref = pages_ref[idx]
            page_idx_comp = pages_comp[idx]
            
            progress_percent = int((idx / num_pages) * 90)
            comparison_progress[task_id]['progress'] = progress_percent
            comparison_progress[task_id]['message'] = f'Comparando página {page_idx_ref + 1} (Ref) vs {page_idx_comp + 1} (Comp)'
            comparison_progress[task_id]['pages_processed'] = idx + 1
            
            t1 = doc1[page_idx_ref].get_text("text")
            t2 = doc2[page_idx_comp].get_text("text")
            
            status, diffs, tipo = compare_content(t1, t2, max_phrase_length=max_phrase_length)
            
            if status == "DIFERENTE" and tipo == "CONTENIDO":
                for diff in diffs:
                    if len(errores_contenido) >= max_errors:
                        break
                    
                    # Generar capturas de pantalla
                    comp = generate_comparison(pdf1_path, pdf2_path, page_idx_ref, page_idx_comp, diff, temp_dir)
                    comp['pagina'] = page_idx_ref + 1
                    comp['pagina_ref'] = page_idx_ref + 1
                    comp['pagina_comp'] = page_idx_comp + 1
                    
                    errores_contenido.append(comp)
                    comparison_progress[task_id]['errors_content'] = len(errores_contenido)
                
                if len(errores_contenido) >= max_errors:
                    break
                    
            elif status == "DIFERENTE" and tipo == "FORMATO":
                errores_formato.append({'pagina': f"Ref:{page_idx_ref + 1} / Comp:{page_idx_comp + 1}"})
                comparison_progress[task_id]['errors_format'] = len(errores_formato)
        
        doc1.close()
        doc2.close()
        
        # Generar documento Word con capturas
        comparison_progress[task_id]['progress'] = 95
        comparison_progress[task_id]['message'] = 'Generando documento Word con capturas...'
        
        doc = Document()
        doc.add_heading('COMPARACIÓN DE PDFs CON CAPTURAS', 0)
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Páginas analizadas: {num_pages}")
        doc.add_paragraph(f"Diferencias de CONTENIDO: {len(errores_contenido)}")
        doc.add_paragraph(f"Diferencias de FORMATO: {len(errores_formato)}")
        
        doc.add_page_break()
        doc.add_heading("DIFERENCIAS DE CONTENIDO", level=1)
        
        if errores_contenido:
            for err in errores_contenido[:100]:
                doc.add_heading(f"Página {err['pagina']}", level=2)
                doc.add_paragraph(f"Palabra/Frase: {err['palabra']}")
                doc.add_paragraph(f"Tipo: {err['tipo']}")
                
                # Agregar tabla con capturas de pantalla
                tbl = doc.add_table(rows=2, cols=2)
                tbl.style = 'Light Grid Accent 1'
                tbl.rows[0].cells[0].text = "PDF REFERENCIA"
                tbl.rows[0].cells[1].text = "PDF A COMPARAR"
                
                # Insertar imágenes si existen
                if os.path.exists(err['img_p']):
                    try:
                        tbl.rows[1].cells[0].paragraphs[0].add_run().add_picture(err['img_p'], width=Inches(3.0))
                    except:
                        pass
                
                if os.path.exists(err['img_s']):
                    try:
                        tbl.rows[1].cells[1].paragraphs[0].add_run().add_picture(err['img_s'], width=Inches(3.0))
                    except:
                        pass
                
                doc.add_paragraph()
        else:
            doc.add_paragraph("No hay diferencias de contenido")
        
        doc.add_page_break()
        doc.add_heading("DIFERENCIAS DE FORMATO", level=1)
        
        if errores_formato:
            for err in errores_formato:
                doc.add_paragraph(f"Página {err['pagina']}: Diferencias de formato detectadas")
        else:
            doc.add_paragraph("No hay diferencias de formato")
        
        output_filename = f"Comparacion_{task_id}.docx"
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        comparison_progress[task_id]['progress'] = 100
        comparison_progress[task_id]['status'] = 'completed'
        comparison_progress[task_id]['message'] = 'Comparación completada'
        comparison_progress[task_id]['result_file'] = output_filename
        
    except Exception as e:
        comparison_progress[task_id]['status'] = 'error'
        comparison_progress[task_id]['message'] = f'Error: {str(e)}'

@app.route('/')
def index():
    return render_template('index.html')


@app.route("/upload-batch", methods=["POST"])
def upload_batch():
    if "pdfs_reference" not in request.files or "pdfs_compare" not in request.files:
        return jsonify({"error": "Faltan archivos"}), 400
    refs = request.files.getlist("pdfs_reference")
    comps = request.files.getlist("pdfs_compare")
    if not refs or not comps:
        return jsonify({"error": "No se seleccionaron archivos"}), 400
    batch_id = str(uuid.uuid4())
    batch_dir = os.path.join(UPLOAD_FOLDER, f"batch_{batch_id}")
    ref_dir = os.path.join(batch_dir, "reference")
    comp_dir = os.path.join(batch_dir, "compare")
    os.makedirs(ref_dir, exist_ok=True)
    os.makedirs(comp_dir, exist_ok=True)
    ref_files = []
    for f in refs:
        if allowed_file(f.filename):
            filename = secure_filename(f.filename)
            f.save(os.path.join(ref_dir, filename))
            ref_files.append(filename)
    comp_files = []
    for f in comps:
        if allowed_file(f.filename):
            filename = secure_filename(f.filename)
            f.save(os.path.join(comp_dir, filename))
            comp_files.append(filename)
    batch_sessions[batch_id] = {"ref_dir": ref_dir, "comp_dir": comp_dir, "ref_files": ref_files, "comp_files": comp_files, "pairs": []}
    return jsonify({"batch_id": batch_id, "ref_files": ref_files, "comp_files": comp_files})

@app.route('/upload', methods=['POST'])
def upload_files():
    """Endpoint para subir archivos PDF (individual o múltiple)"""
    mode = request.form.get('mode', 'single')
    
    if mode == 'single':
        # Modo PDF individual
        if 'pdf_reference' not in request.files or 'pdf_compare' not in request.files:
            return jsonify({'error': 'Faltan archivos'}), 400
        
        pdf_ref = request.files['pdf_reference']
        pdf_comp = request.files['pdf_compare']
        
        if pdf_ref.filename == '' or pdf_comp.filename == '':
            return jsonify({'error': 'No se seleccionaron archivos'}), 400
        
        if not (allowed_file(pdf_ref.filename) and allowed_file(pdf_comp.filename)):
            return jsonify({'error': 'Solo se permiten archivos PDF'}), 400
        
        task_id = str(uuid.uuid4())
        task_dir = os.path.join(UPLOAD_FOLDER, task_id)
        os.makedirs(task_dir, exist_ok=True)
        
        ref_filename = secure_filename(pdf_ref.filename)
        comp_filename = secure_filename(pdf_comp.filename)
        
        ref_path = os.path.join(task_dir, f"reference_{ref_filename}")
        comp_path = os.path.join(task_dir, f"compare_{comp_filename}")
        
        pdf_ref.save(ref_path)
        pdf_comp.save(comp_path)
        
        ref_pages = get_pdf_pages(ref_path)
        comp_pages = get_pdf_pages(comp_path)
        
        return jsonify({
            'task_id': task_id,
            'ref_pages': ref_pages,
            'comp_pages': comp_pages,
            'ref_filename': ref_filename,
            'comp_filename': comp_filename
        })
    
    else:
        # Modo múltiple (carpeta)
        if 'pdfs_reference' not in request.files or 'pdfs_compare' not in request.files:
            return jsonify({'error': 'Faltan archivos'}), 400
        
        refs = request.files.getlist('pdfs_reference')
        comps = request.files.getlist('pdfs_compare')
        
        if not refs or not comps:
            return jsonify({'error': 'No se seleccionaron archivos'}), 400
        
        # Crear diccionarios por nombre de archivo
        refs_dict = {secure_filename(f.filename): f for f in refs if allowed_file(f.filename)}
        comps_dict = {secure_filename(f.filename): f for f in comps if allowed_file(f.filename)}
        
        # Encontrar pares coincidentes
        common_names = set(refs_dict.keys()) & set(comps_dict.keys())
        
        if not common_names:
            return jsonify({'error': 'No se encontraron archivos con nombres coincidentes'}), 400
        
        tasks = []
        batch_id = str(uuid.uuid4())
        
        for filename in sorted(common_names):
            task_id = str(uuid.uuid4())
            task_dir = os.path.join(UPLOAD_FOLDER, task_id)
            os.makedirs(task_dir, exist_ok=True)
            
            ref_path = os.path.join(task_dir, f"reference_{filename}")
            comp_path = os.path.join(task_dir, f"compare_{filename}")
            
            refs_dict[filename].save(ref_path)
            comps_dict[filename].save(comp_path)
            
            tasks.append({
                'task_id': task_id,
                'ref_filename': filename,
                'comp_filename': filename,
                'batch_id': batch_id
            })
        
        return jsonify({
            'batch_id': batch_id,
            'tasks': tasks
        })

@app.route('/compare', methods=['POST'])
def start_comparison():
    """Inicia la comparación de PDFs (individual) con rangos separados"""
    data = request.json
    task_id = data.get('task_id')
    
    # Parámetros separados para cada PDF
    start_page_ref = data.get('start_page_ref', 1)
    end_page_ref = data.get('end_page_ref')
    start_page_comp = data.get('start_page_comp', 1)
    end_page_comp = data.get('end_page_comp')
    custom_pages_ref = data.get('custom_pages_ref')  # Lista de páginas para referencia
    custom_pages_comp = data.get('custom_pages_comp')  # Lista de páginas para comparar
    
    max_errors = data.get('max_errors', 500)
    max_phrase_length = data.get('max_phrase_length', 10)
    
    if not task_id:
        return jsonify({'error': 'Task ID requerido'}), 400
    
    task_dir = os.path.join(UPLOAD_FOLDER, task_id)
    
    if not os.path.exists(task_dir):
        return jsonify({'error': 'Sesión no encontrada'}), 404
    
    files = os.listdir(task_dir)
    ref_file = next((f for f in files if f.startswith('reference_')), None)
    comp_file = next((f for f in files if f.startswith('compare_')), None)
    
    if not ref_file or not comp_file:
        return jsonify({'error': 'Archivos no encontrados'}), 404
    
    ref_path = os.path.join(task_dir, ref_file)
    comp_path = os.path.join(task_dir, comp_file)
    
    result_dir = os.path.join(RESULTS_FOLDER, task_id)
    os.makedirs(result_dir, exist_ok=True)
    
    thread = threading.Thread(
        target=process_comparison_web,
        args=(task_id, ref_path, comp_path, result_dir),
        kwargs={
            'start_page_ref': start_page_ref,
            'end_page_ref': end_page_ref,
            'start_page_comp': start_page_comp,
            'end_page_comp': end_page_comp,
            'custom_pages_ref': custom_pages_ref,
            'custom_pages_comp': custom_pages_comp,
            'max_errors': max_errors,
            'max_phrase_length': max_phrase_length
        }
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({'status': 'started', 'task_id': task_id})

@app.route('/compare-batch', methods=['POST'])
def start_batch_comparison():
    """Inicia comparación por lotes"""
    data = request.json
    task_ids = data.get('tasks', [])
    max_errors = data.get('max_errors', 500)
    max_phrase_length = data.get('max_phrase_length', 10)
    
    if not task_ids:
        return jsonify({'error': 'No hay tareas'}), 400
    
    batch_id = task_ids[0]  # Usar primer task_id como batch_id
    
    batch_progress[batch_id] = {
        'status': 'processing',
        'overall_progress': 0,
        'message': 'Iniciando comparación por lotes...',
        'tasks': []
    }
    
    def process_batch():
        total_tasks = len(task_ids)
        
        for idx, task_id in enumerate(task_ids):
            task_dir = os.path.join(UPLOAD_FOLDER, task_id)
            
            if not os.path.exists(task_dir):
                continue
            
            files = os.listdir(task_dir)
            ref_file = next((f for f in files if f.startswith('reference_')), None)
            comp_file = next((f for f in files if f.startswith('compare_')), None)
            
            if not ref_file or not comp_file:
                continue
            
            ref_path = os.path.join(task_dir, ref_file)
            comp_path = os.path.join(task_dir, comp_file)
            result_dir = os.path.join(RESULTS_FOLDER, task_id)
            os.makedirs(result_dir, exist_ok=True)
            
            batch_progress[batch_id]['message'] = f'Procesando {ref_file}...'
            batch_progress[batch_id]['tasks'].append({
                'filename': ref_file,
                'status': 'processing'
            })
            
            # Procesar comparación
            process_comparison_web(task_id, ref_path, comp_path, result_dir, 
                                 1, None, 1, None, None, None, max_errors, max_phrase_length)
            
            # Actualizar progreso
            batch_progress[batch_id]['tasks'][-1]['status'] = 'completed'
            batch_progress[batch_id]['overall_progress'] = int(((idx + 1) / total_tasks) * 100)
        
        batch_progress[batch_id]['status'] = 'completed'
        batch_progress[batch_id]['message'] = 'Comparación por lotes completada'
    
    thread = threading.Thread(target=process_batch)
    thread.daemon = True
    thread.start()
    
    return jsonify({'status': 'started', 'batch_id': batch_id})

@app.route('/progress/<task_id>')
def get_progress(task_id):
    """Obtiene el progreso de una comparación individual"""
    if task_id not in comparison_progress:
        return jsonify({'error': 'Tarea no encontrada'}), 404
    
    return jsonify(comparison_progress[task_id])

@app.route('/progress-batch/<batch_id>')
def get_batch_progress(batch_id):
    """Obtiene el progreso de una comparación por lotes"""
    if batch_id not in batch_progress:
        return jsonify({'error': 'Lote no encontrado'}), 404
    
    return jsonify(batch_progress[batch_id])

@app.route('/download/<task_id>')
def download_result(task_id):
    """Descarga el resultado de la comparación"""
    if task_id not in comparison_progress:
        return jsonify({'error': 'Tarea no encontrada'}), 404
    
    progress = comparison_progress[task_id]
    
    if progress['status'] != 'completed':
        return jsonify({'error': 'Comparación no completada'}), 400
    
    result_file = progress.get('result_file')
    if not result_file:
        return jsonify({'error': 'Archivo no encontrado'}), 404
    
    file_path = os.path.join(RESULTS_FOLDER, task_id, result_file)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'Archivo no existe'}), 404
    
    return send_file(file_path, as_attachment=True, download_name=result_file)

@app.route('/cancel/<task_id>', methods=['POST'])
def cancel_comparison(task_id):
    """Cancela una comparación en progreso"""
    if task_id in comparison_progress:
        comparison_progress[task_id]['cancelled'] = True
        return jsonify({'status': 'cancelled'})
    return jsonify({'error': 'Tarea no encontrada'}), 404

@app.route('/cleanup/<task_id>', methods=['POST'])
def cleanup_task(task_id):
    """Limpia los archivos de una tarea"""
    try:
        task_dir = os.path.join(UPLOAD_FOLDER, task_id)
        if os.path.exists(task_dir):
            shutil.rmtree(task_dir)
        
        result_dir = os.path.join(RESULTS_FOLDER, task_id)
        if os.path.exists(result_dir):
            shutil.rmtree(result_dir)
        
        if task_id in comparison_progress:
            del comparison_progress[task_id]
        
        return jsonify({'status': 'cleaned'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)  
